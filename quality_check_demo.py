"""
quality_check_demo.py
Self-contained demo:
• Creates two DOCX “contracts”
• Runs a Flask clause checker
• Uses LangChain + local Ollama to call the checker
"""

from flask import Flask, request, jsonify
from docx import Document
import threading, os, requests
from pathlib import Path

# -------------------------------------------------------------------- #
# 0.  Make sure we have sample documents                               #
# -------------------------------------------------------------------- #
os.makedirs("documents", exist_ok=True)


def create_doc(filename, indemnity=True, non_compete=True, law=True):
    doc = Document()
    doc.add_heading("M&A Agreement", level=1)
    doc.add_paragraph("Date: 2024-01-01")
    doc.add_paragraph("Parties: Acme Corp and Target LLC")
    if indemnity:
        doc.add_paragraph("Indemnity: The parties agree to indemnify each other.")
    if non_compete:
        doc.add_paragraph("Non-compete: The seller shall not compete.")
    if law:
        doc.add_paragraph("Governing Law: Delaware applies.")
    doc.save(filename)


create_doc("documents/doc123.docx", indemnity=False)
create_doc("documents/doc124.docx", non_compete=False, law=False)

# -------------------------------------------------------------------- #
# 1.  Tiny Flask API                                                   #
# -------------------------------------------------------------------- #
app = Flask(__name__)

CLAUSES = {
    "indemnity": "indemnity",
    "non-compete": "non-compete",
    "governing-law": "governing law",
}


def check_contract(path: str) -> dict:
    if not os.path.exists(path):
        return {"error": f"File not found: {path}"}

    try:
        doc = Document(path)
        text = "\n".join(p.text.lower() for p in doc.paragraphs)
        missing = [c for c, kw in CLAUSES.items() if kw not in text]
        return {
            "missing_clauses": missing,
            "date_mismatch": "2024" not in text,
            "entity_mismatch": "acme corp" not in text,
        }
    except Exception as e:
        return {"error": str(e)}


@app.route("/mcp/quality", methods=["POST"])
def quality_endpoint():
    return jsonify(check_contract(request.json["path"]))


threading.Thread(target=app.run, kwargs={"port": 5000}, daemon=True).start()

# -------------------------------------------------------------------- #
# 2.  LangChain tool                                                   #
# -------------------------------------------------------------------- #
from pydantic import BaseModel, Field
from langchain.tools import tool


class QualityCheckInput(BaseModel):
    path: str = Field(description="DOCX file name or full path")


@tool(args_schema=QualityCheckInput)
def run_quality_check(path: str) -> str:
    """Run a quality check on a contract and summarize the findings."""
    path = path.strip().strip("'\"")
    p = Path(path)
    if p.suffix.lower() != ".docx":
        p = p.with_suffix(".docx")
    if not p.is_absolute():
        p = Path("documents") / p

    r = requests.post("http://localhost:5000/mcp/quality", json={"path": str(p)})
    result = r.json()

    if "error" in result:
        return f"Error: {result['error']}"

    missing = result["missing_clauses"]
    msg = []

    if missing:
        msg.append(f"Missing clauses: {', '.join(missing)}.")
    else:
        msg.append("All required clauses are present.")

    if result["date_mismatch"]:
        msg.append("Date mismatch found.")
    if result["entity_mismatch"]:
        msg.append("Entity mismatch found.")

    return " ".join(msg)


# -------------------------------------------------------------------- #
# 3.  LLM + ReAct agent (with iteration cap)                           #
# -------------------------------------------------------------------- #
from langchain_ollama import ChatOllama
from langchain.agents import initialize_agent, AgentType

llm = ChatOllama(model="llama3.2", temperature=0)

agent = initialize_agent(
    tools=[run_quality_check],
    llm=llm,
    agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
    verbose=True,
    max_iterations=5,                 # ← stop runaway loops
    early_stopping_method="generate", # ← produce final answer
)

# -------------------------------------------------------------------- #
# 4.  Demo                                                             #
# -------------------------------------------------------------------- #
if __name__ == "__main__":
    print("=== Raw tool sanity-check ===")
    print(run_quality_check.invoke("doc123"))
    print(run_quality_check.invoke("doc124.docx"))

    print("\n=== Agent calls ===")
    res1 = agent.invoke({"input": "Run a quality check on doc123"})
    print(res1["output"])

    res2 = agent.invoke({"input": "Check doc124 for missing clauses"})
    print(res2["output"])
